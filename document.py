

from docx import Document
from docx import *
def make_document(file_name):
    document = Document()
    document.save(file_name)

    document.add_picture('IV_logo.png')
    document.add_paragraph('testts', style='ListBullet')

make_document('test.docx')


